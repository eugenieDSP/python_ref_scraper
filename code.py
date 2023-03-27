import re
import sys
from collections import defaultdict
from docx import Document
from scholarly import scholarly
import bibtexparser
from bibtexparser.bwriter import BibTexWriter
import tkinter as tk
from tkinter import filedialog

# Install required libraries if not installed
# !pip install python-docx scholarly bibtexparser

# Find citations in the text using a regex pattern
def find_citations(doc_text):
    citation_regex = r'\b(\w+(?:\s+and\s+\w+)?(?:\s+et\s+al\.)?)\s*(?:\((\d{4})\)|(\d{4})|\,\s*(\d{4}))'
    matches = re.findall(citation_regex, doc_text)
    citations = [(match[0], match[1] if match[1] else match[2] if match[2] else match[3]) for match in matches]
    return sorted(set(citations), key=lambda x: x[1].lower())

# Fetch paper information from Google Scholar
def fetch_paper_info(authors, year):
    query = f"{authors} {year}"
    search_results = scholarly.search_pubs(query)

    for result in search_results:
        try:
            bib = result.bib
            if 'author' in bib and 'year' in bib and bib['year'] == year:
                return bib
        except Exception as e:
            print(f"Error fetching paper info: {e}")

    return None

# Create a BibTeX entry from the paper information
def create_bibtex_entry(paper_info):
    entry = bibtexparser.customization.author(paper_info)
    entry['ENTRYTYPE'] = 'article'
    entry['ID'] = f"{entry['author'][0]['last']}_{entry['year']}"
    return entry

# Main function
def main(doc_file):
    # Read the .doc file
    document = Document(doc_file)
    doc_text = ' '.join([para.text for para in document.paragraphs])

    # Find citations in the text
    citations = find_citations(doc_text)
    # Create a new BibDatabase
    bib_database = bibtexparser.bibdatabase.BibDatabase()
    bib_database.entries = []

    # Fetch paper information for each citation and add it to the BibDatabase
    for authors, year in citations:
        paper_info = fetch_paper_info(authors, year)
        if paper_info:
            bib_entry = create_bibtex_entry(paper_info)
            bib_database.entries.append(bib_entry)
            
    # Write the BibDatabase to a .bib file
    with open('bibliography.bib', 'w') as bib_file:
        bib_writer = BibTexWriter()
        bib_file.write(bib_writer.write(bib_database))

    print("Bibliography created in 'bibliography.bib' file.")

# Entry point
if __name__ == "__main__":
    # Create a simple Tkinter window
    root = tk.Tk()
    root.withdraw()

    # Open the file explorer and store the selected file's path
    doc_file = filedialog.askopenfilename(title="Select your .docx file", filetypes=(("docx files", "*.docx"),))

    if doc_file:
        main(doc_file)
    else:
        print("No file selected.")
    # Create .doc file with the full bibliography
    bib_doc = Document()

    for entry in bib_database.entries:
        author = " and ".join([f"{a['first']} {a['last']}" for a in entry['author']])
        title = entry.get('title', '')
        year = entry.get('year', '')
        doi = entry.get('doi', '')
        bib_line = f"{author}. ({year}). {title}. DOI: {doi}"
        bib_doc.add_paragraph(bib_line)

    # Save the bibliography to a .docx file
    bib_doc.save('bibliography.docx')
    print("Bibliography created in 'bibliography.docx' file.")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python script.py <doc_file>")
    else:
        doc_file = sys.argv[1]
        main(doc_file)
